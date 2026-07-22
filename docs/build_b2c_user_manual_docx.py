"""Build the B2C user manual as a Word document with embedded screenshots."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "b2c-user-manual.md"
OUTPUT = ROOT / "b2c-user-manual.docx"


IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")


def set_run_font(run, *, size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = "Aptos"
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_inline_markdown(paragraph, text: str) -> None:
    """Render a small subset of inline markdown used in the manual."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Aptos Mono"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(72, 80, 110)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_code_block(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph()
    paragraph.style = "Intense Quote"
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Aptos Mono"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(32, 41, 64)


def add_picture(document: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[Screenshot missing: {caption}]").italic = True
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.4))

    if caption:
        caption_paragraph = document.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption)
        caption_run.italic = True
        caption_run.font.size = Pt(9)
        caption_run.font.color.rgb = RGBColor(95, 105, 128)


def build() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(11)

    code_lines: list[str] = []
    in_code = False

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line == "---":
            continue

        image_match = IMAGE_RE.fullmatch(line)
        if image_match:
            add_picture(document, ROOT / image_match.group("path"), image_match.group("alt"))
            continue

        if line.startswith("# "):
            paragraph = document.add_heading(line[2:], level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if line.startswith("## "):
            document.add_heading(line[3:], level=1)
            continue

        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, line[2:])
            continue

        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, numbered.group(1))
            continue

        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, line)

    if code_lines:
        add_code_block(document, code_lines)

    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build()

"""Build a Word document of the project's most important code snippets.

The goal is not to duplicate the entire source tree. This creates a curated,
submission-friendly document that highlights the unique engineering choices in
Skairova with readable, industry-level comments.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "skairova-important-code-snippets.docx"


SNIPPETS = [
    {
        "title": "1. Gemini Adapter For AI Search",
        "why": (
            "This wrapper keeps the rest of the Flask app independent from the exact Gemini SDK shape. "
            "If the model or SDK changes, only this small boundary needs to change."
        ),
        "code": '''# app.py
# Central model name so upgrades are deliberate and easy to audit.
GEMINI_MODEL_NAME = "gemini-2.5-flash"


class _GeminiModel:
    """Small adapter around Google's GenAI client.

    The rest of the app calls `generate_content(prompt)` and does not need to
    know the SDK-specific syntax. This reduces vendor lock-in and makes future
    model upgrades safer.
    """

    def __init__(self, client, model_name: str):
        self.client = client
        self.model_name = model_name

    def generate_content(self, prompt: str):
        return self.client.models.generate_content(
            model=self.model_name,
            contents=prompt,
        )


try:
    from google import genai

    # If GOOGLE_API_KEY is missing, AI routes fail gracefully instead of
    # crashing the whole flight-search platform at startup.
    model = _GeminiModel(genai.Client(api_key=GOOGLE_API_KEY), GEMINI_MODEL_NAME) if GOOGLE_API_KEY else None
except Exception:
    model = None''',
    },
    {
        "title": "2. Natural-Language Flight Request Parser",
        "why": (
            "This is one of the most important B2C features: users can type a sentence instead of filling a form. "
            "The code combines Gemini extraction with deterministic cleanup so the app gets predictable search parameters."
        ),
        "code": '''# app.py
def parse_ai_flight_request(user_text: str) -> dict | None:
    """Convert a human sentence into structured flight-search parameters."""

    if not model or not user_text:
        return None

    today = date.today().isoformat()
    prompt = f"""
You are a flight search assistant.

Convert the user's request into valid JSON with these fields:
- origin (IATA code or null)
- destination (IATA code or null)
- depart_date (YYYY-MM-DD or null)
- return_date (YYYY-MM-DD or null)
- legs (array of objects: origin, destination, depart_date) for multi-city requests
- trip_type (oneway, roundtrip, multicity, or null)
- passengers (integer, default 1)
- cabin (ECONOMY, PREMIUM_ECONOMY, BUSINESS, FIRST)
- nonstop (true/false)
- max_price (number or null)
- sort (cheapest, fastest, recommended)

Rules:
- Use null if information is missing.
- If a city is mentioned, infer the main airport.
- Dates must be ISO format.
- Today is {today}.
- Only return JSON.

User request:
\"\"\"{user_text}\"\"\"
"""

    try:
        response = model.generate_content(prompt)
        text = (getattr(response, "text", "") or "").strip()

        # Gemini may return fenced JSON in development. Strip that safely.
        if text.startswith("```"):
            text = text.strip().strip("`").strip()
            if text.lower().startswith("json"):
                text = text[4:].strip()

        parsed = json.loads(text)

        # Normalize AI output before it reaches Duffel or validation logic.
        parsed["origin"] = _normalize_airport_input(parsed.get("origin")) if parsed.get("origin") else None
        parsed["destination"] = _normalize_airport_input(parsed.get("destination")) if parsed.get("destination") else None
        parsed["passengers"] = int(parsed.get("passengers") or 1)
        parsed["cabin"] = parsed.get("cabin") or "ECONOMY"
        parsed["nonstop"] = bool(parsed.get("nonstop") or False)
        parsed["sort"] = parsed.get("sort") or "recommended"

        # Avoid a common UX bug: if no return date exists, default to one-way.
        fallback_trip_type = "roundtrip" if parsed.get("return_date") else "oneway"
        parsed["trip_type"] = _coerce_trip_type(parsed.get("trip_type"), fallback=fallback_trip_type)
        parsed.setdefault("raw_text", user_text)

        return parsed
    except Exception as e:
        print("JSON PARSE ERROR:", e)
        return None''',
    },
    {
        "title": "3. Flexible Month And Holiday Date Intelligence",
        "why": (
            "This is what makes the AI search feel travel-aware rather than just text-to-JSON. "
            "It can understand phrases like 'next month', 'one week', or 'Thanksgiving trip' and turn them into useful search behavior."
        ),
        "code": '''# app.py
def _looks_like_ai_flex_request(user_text: str, parsed: dict[str, Any] | None = None) -> bool:
    """Decide whether an AI request should trigger the cheapest-week scanner."""

    txt = (user_text or "").strip().lower()
    if not txt:
        return False

    # Exact dates should remain exact searches, not flexible-month searches.
    if _user_text_has_explicit_day_precision(txt):
        return False

    has_month = _extract_ai_flex_month(txt) is not None
    has_length = _extract_ai_trip_length_days(txt) is not None
    trip_type = _extract_ai_trip_type(txt, parsed)

    if trip_type == "multicity":
        return False

    has_relative_month_phrase = bool(re.search(r"\\bnext month\\b|\\b(?:this|current) month\\b", txt))
    flex_words = any(w in txt for w in ["cheapest week", "cheapest ", "best ", "flexible", "any time in", "during "])

    return has_month and (
        has_length
        or flex_words
        or trip_type == "oneway"
        or has_relative_month_phrase
    )


def _infer_holiday_season_round_trip(
    user_text: str,
    *,
    anchor: date,
    parsed: dict[str, Any] | None = None,
) -> tuple[str, str] | None:
    """Map common travel phrases to editable default round-trip windows.

    Example: "Thanksgiving trip to LA" becomes a sensible Wednesday-Sunday
    holiday window without forcing the user to know exact calendar dates.
    """

    txt = (user_text or "").strip().lower()
    if not txt or re.search(r"\\bone[\\s-]?way\\b", txt):
        return None

    def tg_rt(y: int) -> tuple[date, date]:
        thanksgiving = _us_thanksgiving(y)
        return thanksgiving - timedelta(days=1), thanksgiving + timedelta(days=3)

    def xmas_rt(y: int) -> tuple[date, date]:
        return date(y, 12, 22), date(y, 12, 27)

    if re.search(r"\\bthanksgiving\\b|\\bblack\\s+friday\\b", txt):
        return _next_round_trip_window(anchor, tg_rt)

    if re.search(r"\\b(christmas|xmas)\\b|\\bholiday\\s+season\\b", txt):
        return _next_round_trip_window(anchor, xmas_rt)

    return None''',
    },
    {
        "title": "4. Reloadable Results After POST Search",
        "why": (
            "A normal POST search disappears on browser refresh. This code stores rendered results behind a safe token, "
            "then redirects to a shareable/reloadable GET URL."
        ),
        "code": '''# app.py
_RESULTS_RELOAD_TOKEN_RE = re.compile(r"^[A-Za-z0-9_-]{16,96}$")


def _new_results_reload_token() -> str:
    """Generate an opaque token for one rendered results page."""
    return os.urandom(16).hex()


def _store_results_reload_html(html: str, *, token: str | None = None) -> str:
    """Cache completed results HTML so refreshing `/results/<token>` works."""

    reload_token = _coerce_results_reload_token(token) or _new_results_reload_token()
    RESULTS_RELOAD_CACHE.set(
        reload_token,
        {
            "html": html,
            "created_at": int(time.time()),
        },
    )
    return reload_token


@app.route("/results/<token>", methods=["GET"])
def results_reload(token: str):
    """Serve cached results through a safe GET route."""

    reload_token = _coerce_results_reload_token(token)
    payload = RESULTS_RELOAD_CACHE.get(reload_token) if reload_token else None

    if isinstance(payload, Mapping):
        html = str(payload.get("html") or "")
        if html:
            return Response(html, mimetype="text/html")

    return render_template(
        "results.html",
        query={},
        flights=[],
        error="These cached results have expired. Please run the search again to refresh live fares.",
        minutes_to_hm=minutes_to_hm,
        fmt_dt=fmt_dt,
    ), 410


@app.after_request
def _redirect_post_search_results_to_reload_url(response: Response):
    """Convert completed POST search responses into reloadable GET pages."""

    if (
        request.endpoint == "search"
        and request.method == "POST"
        and response.status_code == 200
        and response.mimetype == "text/html"
        and not response.is_streamed
    ):
        html = response.get_data(as_text=True)
        if html:
            reload_token = _store_results_reload_html(html)
            return redirect(url_for("results_reload", token=reload_token), code=303)

    return response''',
    },
    {
        "title": "5. Resilient Duffel API Client",
        "why": (
            "This is the integration boundary to live flight inventory and booking. "
            "It centralizes authentication, retry behavior, timeout handling, and user-safe error messages."
        ),
        "code": '''# app.py
def _build_session(*, retry_total: int = 3, backoff_factor: float = 0.35) -> requests.Session:
    """Create a pooled HTTP session with retry behavior for transient failures."""

    session = requests.Session()
    retry = Retry(
        total=retry_total,
        connect=retry_total,
        read=retry_total,
        backoff_factor=backoff_factor,
        status_forcelist=(429, 500, 502, 503, 504),
        allowed_methods=frozenset(["GET", "POST"]),
        raise_on_status=False,
    )
    adapter = HTTPAdapter(
        max_retries=retry,
        pool_connections=HTTP_POOL_CONNECTIONS,
        pool_maxsize=HTTP_POOL_MAXSIZE,
    )
    session.mount("https://", adapter)
    session.mount("http://", adapter)
    session.headers.update({"Accept": "application/json"})
    return session


class DuffelClient:
    """Typed boundary around Duffel's HTTP API."""

    def __init__(self):
        self.session = _build_session()
        self.fast_session = _build_session(retry_total=0, backoff_factor=0.0)

    def _headers(self) -> dict[str, str]:
        return {
            "Authorization": f"Bearer {DUFFEL_ACCESS_TOKEN}",
            "Duffel-Version": DUFFEL_VERSION,
            "Accept": "application/json",
            "Content-Type": "application/json",
        }

    def _request(
        self,
        method: str,
        path: str,
        *,
        params: dict[str, Any] | None = None,
        json_body: dict[str, Any] | None = None,
        timeout: float = DUFFEL_HTTP_TIMEOUT,
        fast: bool = False,
    ) -> requests.Response:
        session = self.fast_session if fast else self.session
        return session.request(
            method=method,
            url=f"{DUFFEL_BASE}{path}",
            params=params,
            json=json_body,
            headers=self._headers(),
            timeout=timeout,
        )''',
    },
    {
        "title": "6. Booking Creation With Validation, Ancillaries, Tracking, And Email",
        "why": (
            "This is the central checkout transaction. It validates the traveler form, checks offer expiry, "
            "adds selected services, creates the Duffel order, tracks the funnel, and sends itinerary emails."
        ),
        "code": '''# app.py
@app.route("/checkout/<offer_id>/details", methods=["GET", "POST"])
def checkout_details(offer_id: str):
    """Render checkout and create the booking on POST."""

    offer = DUFF.get_offer(offer_id, return_available_services=True)
    seat_maps, payment_config = _load_checkout_sidecars(offer)
    ancillaries_payload = extract_ancillaries_payload(
        request.form if request.method == "POST" else request.args
    )

    travelers = build_traveler_forms(offer, request.form if request.method == "POST" else None)
    offer_summary = build_checkout_summary(
        offer,
        seat_maps=seat_maps,
        ancillaries_payload=ancillaries_payload,
    )
    checkout_model = build_checkout_page_model(
        offer,
        travelers=travelers,
        seat_maps=seat_maps,
        ancillaries_payload=ancillaries_payload,
        payment_config=payment_config,
        duffel_env=DUFFEL_ENV,
    )

    if request.method == "POST":
        if offer_has_expired(offer):
            return render_template(
                "checkout.html",
                offer_summary=offer_summary,
                travelers=travelers,
                checkout_model=checkout_model,
                booking_error="This offer has expired. Please choose a fresh option.",
                booking_enabled=False,
                duffel_env=DUFFEL_ENV,
            ), 410

        passengers_payload, travelers, errors = validate_checkout_form(offer, request.form)
        if errors:
            return render_template(
                "checkout.html",
                offer_summary=offer_summary,
                travelers=travelers,
                checkout_model=checkout_model,
                errors=errors,
                booking_error=errors.get("form", ""),
                booking_enabled=True,
                duffel_env=DUFFEL_ENV,
            ), 400

        selected_services = selected_services_from_payload(ancillaries_payload)
        order_services = normalize_create_order_services(selected_services)
        total_amount = calculate_total_amount(offer, ancillaries_payload, seat_maps=seat_maps)

        order = DUFF.create_order(
            offer_id=(offer.get("id") or offer_id).strip(),
            passengers=passengers_payload,
            services=order_services or None,
            total_amount=str(total_amount or offer_summary.get("total_amount") or "0.00"),
            total_currency=str(offer_summary.get("currency") or "USD"),
        )

        order_id = str(order.get("id") or "").strip()
        if order_id:
            RECENT_ORDER_CACHE.set(order_id, order)
            _capture_booking_email_links(order=order, passengers_payload=passengers_payload)
            _track_booking_completed_event(order, offer=offer)
            _send_itinerary_emails_after_booking(order=order, passengers_payload=passengers_payload)
            return redirect(url_for("booking_confirmation", order_id=order_id))

    return render_template(
        "checkout.html",
        offer_summary=offer_summary,
        travelers=travelers,
        checkout_model=checkout_model,
        booking_enabled=True,
        duffel_env=DUFFEL_ENV,
    )''',
    },
    {
        "title": "7. Baggage Allowance Normalization",
        "why": (
            "Duffel returns baggage per segment and per passenger. This code turns nested provider data into a clean customer-facing summary."
        ),
        "code": '''# duffel_booking.py
def _build_baggage_summary(offer: Mapping[str, Any]) -> list[dict[str, str]]:
    """Summarize included baggage from the first flown segment.

    Duffel reports bag allowance per segment/passenger. The first segment is the
    clearest checkout summary proxy; details can vary later by airline/segment.
    """

    first_slice = (offer.get("slices") or [])[0] if offer.get("slices") else {}
    first_segment = ((first_slice or {}).get("segments") or [{}])[0]
    segment_passengers = (first_segment or {}).get("passengers") or []

    if not segment_passengers:
        return [{
            "type": "unknown",
            "label": "Baggage",
            "quantity_label": "Airline policy",
            "weight_label": "not provided",
        }]

    by_passenger: dict[str, dict[str, int]] = {}
    weights: dict[str, str] = {}

    for index, passenger in enumerate(segment_passengers):
        passenger_id = str((passenger or {}).get("passenger_id") or index)
        by_passenger.setdefault(passenger_id, {})

        for bag in (passenger or {}).get("baggages") or []:
            bag_type = str((bag or {}).get("type") or "").strip().lower()

            # Duffel may use "cabin"; the UI uses the clearer "carry_on".
            if bag_type == "cabin":
                bag_type = "carry_on"
            if bag_type not in {"carry_on", "checked"}:
                continue

            quantity = int((bag or {}).get("quantity") or 0)
            by_passenger[passenger_id][bag_type] = (
                by_passenger[passenger_id].get(bag_type, 0) + max(0, quantity)
            )
            weights.setdefault(bag_type, _bag_weight_label(bag))

    output: list[dict[str, str]] = []
    labels = {"carry_on": "Carry-on", "checked": "Checked bag"}

    for bag_type in ("carry_on", "checked"):
        quantities = [allowances.get(bag_type, 0) for allowances in by_passenger.values()]
        total_quantity = sum(quantities)
        if total_quantity <= 0:
            continue

        quantity_label = (
            f"{quantities[0]} per traveler"
            if len(set(quantities)) == 1
            else f"{total_quantity} total for {len(by_passenger)} travelers"
        )

        output.append({
            "type": bag_type,
            "label": labels[bag_type],
            "quantity_label": quantity_label,
            "weight_label": weights.get(bag_type, "weight not provided"),
        })

    return output or [{
        "type": "unknown",
        "label": "Baggage",
        "quantity_label": "Airline policy",
        "weight_label": "not provided",
    }]''',
    },
    {
        "title": "8. Page-Aware AI Chat Assistant",
        "why": (
            "The chatbot is not generic. It changes its instructions based on whether the user is on results, review, checkout, confirmation, or manage-booking pages."
        ),
        "code": '''# app.py
def _build_ai_chat_system(context: dict) -> str:
    """Create a page-aware system prompt for the embedded Skairova assistant."""

    page_type = str(context.get("page_type") or "results").strip()
    airline = str(context.get("airline_name") or context.get("airline") or "").strip()
    currency = str(context.get("currency") or "USD").strip()
    total = str(context.get("total_amount") or "").strip()
    route = str(context.get("route_summary") or "").strip()
    passengers = context.get("passengers") or context.get("traveler_count") or 1

    base = (
        "You are Skairova AI, a knowledgeable and friendly flight booking expert. "
        "Only discuss topics directly related to travel, flights, airports, airlines, baggage, and logistics. "
        "Keep responses concise and genuinely useful."
    )

    if page_type == "checkout":
        slices = context.get("slices") or []
        itinerary = "\\n".join(
            f"  • {sl.get('label', '')}: departs {sl.get('depart_label', '')}, "
            f"{sl.get('duration_label', '')}, {sl.get('stops_label', '')}"
            for sl in slices[:4]
        )
        return (
            f"{base}\\n\\n"
            "CURRENT PAGE: Checkout — the user is about to confirm and pay.\\n"
            f"SELECTED FLIGHT: {airline or 'Selected airline'}"
            f"{' — ' + route if route else ''}\\n"
            f"ITINERARY:\\n{itinerary}\\n"
            f"TOTAL: {currency} {total}\\n"
            f"TRAVELERS: {passengers}\\n\\n"
            "Help the user feel confident before payment. "
            "Do not invent exact baggage fees or policies."
        )

    if page_type == "confirmation":
        booking_ref = str(context.get("booking_reference") or "").strip()
        return (
            f"{base}\\n\\n"
            "CURRENT PAGE: Booking confirmation — the booking is confirmed.\\n"
            f"BOOKING REFERENCE: {booking_ref or '(unknown)'}\\n"
            f"AIRLINE: {airline or 'Unknown'}"
            f"{' — ' + route if route else ''}\\n"
            "Focus on airport arrival time, online check-in, baggage, and next steps."
        )

    return (
        f"{base}\\n\\n"
        "CURRENT PAGE: Flight search results. "
        "Use only the flights and prices currently provided in page context. "
        "Never invent flight data, prices, or schedules."
    )''',
    },
    {
        "title": "9. AI Chat Endpoint With Context And Conversation History",
        "why": (
            "This endpoint is the bridge between the browser chatbot and Gemini. "
            "It limits input size, carries recent history, and injects page context for safer answers."
        ),
        "code": '''# app.py
@app.route("/api/ai/chat", methods=["POST"])
def ai_chat():
    """Handle chat messages from the floating AI assistant."""

    if not model:
        return jsonify({"error": "AI not configured"}), 503

    data = request.get_json(silent=True) or {}
    user_message = str(data.get("message", "")).strip()[:500]
    if not user_message:
        return jsonify({"error": "Empty message"}), 400

    context = data.get("context", {})
    history = data.get("history", [])
    system_prompt = _build_ai_chat_system(context)

    # Keep recent history only. This preserves conversational continuity while
    # preventing unbounded prompt growth.
    convo_parts = [f"System: {system_prompt}\\n"]
    for h in history[-8:]:
        role = h.get("role", "user")
        msg = str(h.get("content", "")).strip()[:300]
        convo_parts.append(f"{'User' if role == 'user' else 'Assistant'}: {msg}")

    convo_parts.append(f"User: {user_message}")
    convo_parts.append("Assistant:")
    full_prompt = "\\n".join(convo_parts)

    result = model.generate_content(full_prompt)
    text = str(getattr(result, "text", "") or "").strip()
    if not text:
        return jsonify({"error": "No response"}), 500

    return jsonify({"reply": text})''',
    },
    {
        "title": "10. Calendar Itinerary Export",
        "why": (
            "This turns booked flights into a standard `.ics` calendar file, allowing users to add each flight segment to Apple Calendar, Google Calendar, Outlook, and other calendar apps."
        ),
        "code": '''# app.py
def _ics_dt(iso: str | None) -> str:
    """Convert an ISO datetime into UTC iCalendar format."""

    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    except Exception:
        return ""


def _render_ics(summary: dict) -> str:
    """Create one calendar event per flown segment."""

    ref = summary.get("booking_reference") or "UNKNOWN"
    airline = summary.get("airline_name") or "Airline"
    passengers = summary.get("passenger_names") or []
    pax_str = ", ".join(passengers) if passengers else "Traveler"

    events = []
    for sl in summary.get("slices") or []:
        for seg in sl.get("segments") or []:
            dtstart = _ics_dt(seg.get("depart_iso"))
            dtend = _ics_dt(seg.get("arrive_iso"))
            if not dtstart or not dtend:
                continue

            flight_number = seg.get("flight_number") or airline
            origin = seg.get("origin_code") or ""
            destination = seg.get("destination_code") or ""
            origin_city = seg.get("origin_city") or origin
            destination_city = seg.get("destination_city") or destination

            events.append(
                "BEGIN:VEVENT\\r\\n"
                f"UID:{uuid.uuid4()}\\r\\n"
                f"DTSTART:{dtstart}\\r\\n"
                f"DTEND:{dtend}\\r\\n"
                f"SUMMARY:{flight_number}: {origin_city} → {destination_city}\\r\\n"
                f"DESCRIPTION:Booking reference: {ref}\\\\nPassengers: {pax_str}\\r\\n"
                f"LOCATION:{origin_city} ({origin})\\r\\n"
                "END:VEVENT\\r\\n"
            )

    return (
        "BEGIN:VCALENDAR\\r\\n"
        "VERSION:2.0\\r\\n"
        "PRODID:-//Skairova//Flight Booking//EN\\r\\n"
        "CALSCALE:GREGORIAN\\r\\n"
        "METHOD:PUBLISH\\r\\n"
        f"{''.join(events)}"
        "END:VCALENDAR\\r\\n"
    )''',
    },
]


def format_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)


def add_code(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(12)
    paragraph.paragraph_format.left_indent = Inches(0.15)

    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor(23, 37, 58)


def build() -> None:
    document = Document()
    format_doc(document)

    title = document.add_heading("Skairova: Important Code Snippets", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A curated set of the most unique implementation pieces in the B2C flight-booking platform."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(91, 104, 130)

    intro = document.add_paragraph()
    intro.add_run(
        "This document intentionally includes only the highest-value snippets: AI parsing, flexible date intelligence, "
        "reloadable results, Duffel integration, checkout booking, baggage normalization, contextual chatbot logic, "
        "and itinerary export. Comments have been added or tightened for readability."
    )

    for snippet in SNIPPETS:
        document.add_heading(snippet["title"], level=1)
        why = document.add_paragraph()
        label = why.add_run("Why it matters: ")
        label.bold = True
        why.add_run(snippet["why"])
        add_code(document, snippet["code"])

    document.save(OUTPUT)
    print(f"wrote {OUTPUT}")


if __name__ == "__main__":
    build()
